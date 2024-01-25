import tkinter as tk
from tkinter import filedialog
import csv
import re
import subprocess
import os
from urllib.parse import urlparse
from bs4 import BeautifulSoup
from readability import Document
from docx import Document as WordDocument
from docx.shared import Inches
from docx.shared import Pt
from PIL import Image
import base64
import appscript
import mactypes
 
def generate_filename(url):
    parsed_url = urlparse(url)
    fragments = parsed_url.fragment or None
    subdirectory = parsed_url.path.split('/')[1] if len(parsed_url.path.split('/')) > 1 else 'no-subdirectory'
    second_level_domain = parsed_url.netloc.split('.')[-2] if len(parsed_url.netloc.split('.')) > 1 else 'no-second-level-domain'
    top_level_domain = parsed_url.netloc.split('.')[-1] if len(parsed_url.netloc.split('.')) > 0 else 'no-top-level-domain'

    if not fragments:
        fragments = parsed_url.path.split('/')[2] if len(parsed_url.path.split('/')) > 2 else 'no-fragment'

    # Replace hyphens in fragments with spaces
    fragments = fragments.replace('-', ' ')

    filename = f"{subdirectory}-{fragments}-{second_level_domain}.{top_level_domain}"
    filename = re.sub(r'[^\w\s.-]', '', filename)
    return filename

def extract_main_content(html_content):
    doc = Document(html_content)
    main_content = doc.summary(html_partial=True)
    return main_content

def extract_images(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    image_elements = soup.find_all('img')
    return [img['src'] for img in image_elements]

def save_html(url, save_path):
    filename = generate_filename(url)
    command = f'docker run --rm -v "{os.path.abspath(save_path)}":/output singlefile {url}'
    html_content = subprocess.check_output(command, shell=True, text=True)

    # Extract main article content
    main_content = extract_main_content(html_content)

    # Extract images
    images = extract_images(html_content)

    # Set max-width style rule
    max_width_style = "max-width: 500px;"

    # Add max-width style to body tag
    body_tag = f"<body style='{max_width_style}'>"

    # Create Word document
    doc = WordDocument()
    doc.add_paragraph(main_content)

    # Add image placeholders
    for img_src in images:
        doc.add_paragraph(f"Image Placeholder: {img_src}")

    filepath = f'{filename}.docx'
    doc.save(filepath)

    # Add webpage URL to comments using appscript
    file_alias = mactypes.Alias(os.path.abspath(filepath))
    file = appscript.app('Finder').items[file_alias]
    file.comment.set(f"Website: {url}")

    print(f"Processed: {filepath}")

def process_csv(csv_file, save_path):
    with open(csv_file, 'r') as file:
        reader = csv.reader(file)
        total_files = sum(1 for _ in reader)
        file.seek(0)  # Reset file pointer

        print(f"Total number of files: {total_files}")

        for index, row in enumerate(reader, start=1):
            url = row[0]
            print(f"\nProcessing file {index}/{total_files}")
            save_html(url, save_path)

def browse_file():
    file_path = filedialog.askopenfilename(filetypes=[("CSV files", "*.csv")])
    if file_path:
        save_path = filedialog.askdirectory(initialdir="~/Downloads")
        if save_path:
            os.chdir(save_path)
            process_csv(file_path, save_path)
            status_label.config(text="Word documents saved successfully.")
        else:
            status_label.config(text="Operation canceled.")
    else:
        status_label.config(text="Operation canceled.")

# Create the main window
root = tk.Tk()
root.title("URL to Word Converter")

# Create and pack widgets
browse_button = tk.Button(root, text="Browse CSV File", command=browse_file)
browse_button.pack(pady=10)

status_label = tk.Label(root, text="")
status_label.pack(pady=10)

# Run the main loop
root.mainloop()
