import tkinter as tk
from tkinter import filedialog
import csv
import re
import subprocess
import os
from urllib.parse import urlparse
from bs4 import BeautifulSoup
from readability import Document
import appscript
import mactypes
from docx import Document as WordDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import base64
from io import BytesIO
from PIL import Image

# Set to keep track of processed URLs
processed_urls = set()

def is_duplicate(url):
    if url in processed_urls:
        print(f"Skipping duplicate URL: {url}")
        return True
    return False

def generate_filename(url):
    parsed_url = urlparse(url)
    fragments = parsed_url.fragment or None
    subdirectory = parsed_url.path.split('/')[1] if len(parsed_url.path.split('/')) > 1 else 'no-subdirectory'
    second_level_domain = parsed_url.netloc.split('.')[-2] if len(parsed_url.netloc.split('.')) > 1 else 'no-second-level-domain'
    top_level_domain = parsed_url.netloc.split('.')[-1] if len(parsed_url.netloc.split('.')) > 0 else 'no-top-level-domain'

    if not fragments:
        fragments = parsed_url.path.split('/')[2] if len(parsed_url.path.split('/')) > 2 else 'no-fragment'

    # Replace hyphens in fragments with spaces
    fragments = fragments.replace('-', ' ')

    filename = f"{subdirectory}-{fragments}-{second_level_domain}.{top_level_domain}"
    filename = re.sub(r'[^\w\s.-]', '', filename)
    return filename

def extract_main_content(html_content):
    doc = Document(html_content)
    main_content = doc.summary(html_partial=True)
    return main_content

def extract_styles(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    style_elements = soup.find_all('style')
    return '\n'.join([str(style) for style in style_elements])

def extract_images(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    image_elements = soup.find_all('img')
    return '\n'.join([str(img) for img in image_elements])

def remove_iframes(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    for iframe in soup.find_all('iframe'):
        iframe.decompose()  # Remove the iframe element
    return str(soup)

def save_html(url, save_path):
    if is_duplicate(url):
        return

    filename = generate_filename(url)
    command = f'docker run --rm -v "{os.path.abspath(save_path)}":/output singlefile {url}'
    html_content = subprocess.check_output(command, shell=True, text=True)

    # Clean HTML content using readability-lxml
    cleaned_content = extract_main_content(html_content)

    # Remove iframes from cleaned content
    cleaned_content_no_iframes = remove_iframes(cleaned_content)

    # Set max-width style rule
    max_width_style = "max-width: 500px;"

    # Add max-width style to body tag
    body_tag = f"<body style='{max_width_style}'>"

    filepath = f'{filename}.html'

    with open(filepath, 'w', encoding='utf-8') as file:
        file.write(f"<html><head>{extract_styles(html_content)}</head>{body_tag}{cleaned_content_no_iframes}{extract_images(html_content)}</body></html>")

    # Add webpage URL to comments using appscript
    file_alias = mactypes.Alias(os.path.abspath(filepath))
    file = appscript.app('Finder').items[file_alias]
    file.comment.set(f"Website: {url}")

    # Add processed URL to set
    processed_urls.add(url)

    print(f"Processed: {filepath}")

def convert_base64_to_rgb(base64_data):
    # Decode base64 data and convert to RGB
    image_data = base64.b64decode(base64_data)
    image = Image.open(BytesIO(image_data))
    image_rgb = image.convert("RGB")
    return image_rgb

def save_image_as_jpeg(image_rgb, output_path):
    # Save the image as JPEG
    image_rgb.save(output_path, format="JPEG")

def handle_list_item(paragraph, element):
    # Handle list items with bullet points
    run = paragraph.add_run('• ')
    font = run.font
    font.size = Pt(18)  # Font size 18
    paragraph.add_run(element.get_text())

def handle_table(doc, element):
    # Handle tables
    table = doc.add_table(rows=len(element.find_all('tr')), cols=max(len(row.find_all(['td', 'th'])) for row in element.find_all('tr')))
    for row_index, row in enumerate(element.find_all('tr')):
        for col_index, cell in enumerate(row.find_all(['td', 'th'])):
            table.cell(row_index, col_index).text = cell.get_text()

def handle_figure(doc, element, image_counter):
    # Handle figure elements
    if element.find('img') and element.find('img').get('src').startswith('data:image'):
        base64_data = element.find('img')['src'].split(',', 1)[1]
        image_rgb = convert_base64_to_rgb(base64_data)
        temp_image_path = f"temp_image_{image_counter}.jpeg"
        save_image_as_jpeg(image_rgb, temp_image_path)
        doc.add_picture(temp_image_path, width=Inches(5))
        os.remove(temp_image_path)

def handle_math(doc, element):
    # Handle math elements
    run = doc.add_paragraph().add_run(element.get_text())
    font = run.font
    font.size = Pt(18)  # Font size 18

def html_to_word(input_path, output_path):
    # Load HTML content
    with open(input_path, 'r', encoding='utf-8') as html_file:
        html_content = html_file.read()

    # Create a Word document
    doc = WordDocument()

    # Set default font for the entire document
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(18)

    # Parse HTML content
    soup = BeautifulSoup(html_content, 'html.parser')

    # Mapping HTML headings to Word styles
    heading_styles = {
        'h1': 'Heading1',
        'h2': 'Heading2',
        'h3': 'Heading3',
        'h4': 'Heading4',
        'h5': 'Heading5',
        'h6': 'Heading6',
    }

    # Counter for image placeholders
    image_counter = 1

    # Initialize last_line
    last_line = ""

    # Iterate through HTML tags and convert to Word format
    for element in soup.descendants:
        if element.name == 'p':
            paragraph = doc.add_paragraph()

            # Set paragraph alignment if specified in HTML
            if element.get('style'):
                styles = element['style'].split(';')
                for style in styles:
                    if 'text-align' in style:
                        alignment = style.split(':')[-1].strip()
                        if alignment == 'left':
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        elif alignment == 'center':
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        elif alignment == 'right':
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            # Add text to the paragraph
            run = paragraph.add_run(element.get_text())
            font = run.font
            font.size = Pt(18)  # Font size 18

            # Update last_line
            last_line = element.get_text()

        elif element.name in heading_styles:
            # Handle HTML headings
            heading_level = int(element.name[1])
            heading_style = heading_styles[element.name]
            heading = doc.add_paragraph(element.get_text(), style=heading_style)
            heading.style.font.bold = True  # Make heading text bold
            heading.style.font.size = Pt(2 * (18 + 2 * (1 - heading_level)))  # Double and adjust font size based on heading level

            # Update last_line
            last_line = element.get_text()

        elif element.name in ['strong', 'b']:
            # Handle bold text
            bold_text = element.get_text()

            # Check if the last line is the same and delete it
            if last_line.strip() == bold_text.strip():
                # Delete the last paragraph
                doc.paragraphs[-1].clear()

            run = doc.add_paragraph().add_run(bold_text)
            font = run.font
            font.bold = True
            font.size = Pt(18)  # Font size 18

            # Update last_line
            last_line = bold_text

        elif element.name == 'img' and element.get('src') and element.get('src').startswith('data:image'):
            # Handle image elements with base64 data
            base64_data = element['src'].split(',', 1)[1]  # Extract base64 data
            image_rgb = convert_base64_to_rgb(base64_data)

            # Save the image as JPEG
            temp_image_path = f"temp_image_{image_counter}.jpeg"
            save_image_as_jpeg(image_rgb, temp_image_path)

            # Replace the image placeholder with the actual image
            doc.add_picture(temp_image_path, width=Inches(5))
            
            os.remove(temp_image_path)  # Remove the temporary image file

            # Increment the image counter
            image_counter += 1

        elif element.name == 'li':
            # Handle list items
            handle_list_item(doc.add_paragraph(), element)

        elif element.name == 'table':
            # Handle tables
            handle_table(doc, element)

        elif element.name == 'figure':
            # Handle figures
            handle_figure(doc, element, image_counter)

        elif element.name == 'math':
            # Handle math elements
            handle_math(doc, element)

    # Save the Word document
    doc.save(output_path)

    print(f"Conversion successful. Word document saved at: {output_path}")

def process_csv(csv_file, save_path):
    with open(csv_file, 'r') as file:
        reader = csv.reader(file)
        total_files = sum(1 for _ in reader)
        file.seek(0)  # Reset file pointer

        print(f"Total number of files: {total_files}")

        for index, row in enumerate(reader, start=1):
            url = row[0]
            print(f"\nProcessing file {index}/{total_files}")
            save_html(url, save_path)

            # Construct the path for the generated HTML file
            html_file_path = os.path.join(save_path, generate_filename(url) + ".html")

            # Construct the path for the Word output file
            word_output_path = os.path.join(save_path, generate_filename(url) + "_output.docx")

            # Convert the generated HTML to Word
            html_to_word(html_file_path, word_output_path)

# GUI
def browse_file():
    file_path = filedialog.askopenfilename(filetypes=[("CSV files", "*.csv")])
    if file_path:
        save_path = filedialog.askdirectory(initialdir="~/Downloads")
        if save_path:
            os.chdir(save_path)
            processed_urls.clear()  # Clear the set when processing a new CSV file
            process_csv(file_path, save_path)
            status_label.config(text="HTML and Word files saved successfully.")
        else:
            status_label.config(text="Operation canceled.")
    else:
        status_label.config(text="Operation canceled.")

# Create the main window
root = tk.Tk()
root.title("URL to HTML and Word Converter")

# Create and pack widgets
browse_button = tk.Button(root, text="Browse CSV File", command=browse_file)
browse_button.pack(pady=10)

status_label = tk.Label(root, text="")
status_label.pack(pady=10)

# Run the main loop
root.mainloop()
